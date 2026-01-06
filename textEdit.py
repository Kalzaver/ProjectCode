# import sys
import json
from docx import Document


def mass_replace(json_path, copy_file_path):
    print("\n replacement started")
    with open(json_path, 'r', encoding='utf-8') as f:
        replacements = json.load(f)

    doc = Document(copy_file_path)

    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            for run in p.runs:
                for key, value in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))

    replace_in_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

    doc.save(copy_file_path)
    print("replacement finished")

if __name__ == "__main__":
    u_orig_file_path = 'C:/qt_projects/Python_hybrid_test/file_editing/test_2.docx'
    u_json_path = 'C:/qt_projects/Python_hybrid_test/file_editing/data.json'
    u_copy_file_path = 'C:/qt_projects/Python_hybrid_test/file_editing/copy_test3.docx'

    mass_replace(u_json_path, u_copy_file_path)
